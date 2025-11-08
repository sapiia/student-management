#!/usr/bin/env python3
"""
Student Management System - User Manual Generator
Generates a comprehensive Word document user manual for the SMS application.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

def create_user_manual():
    """Create a comprehensive user manual document."""

    # Create document
    doc = Document()

    # Add title page
    title = doc.add_heading('Student Management System', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_heading('User Manual and Documentation', 1)
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add version and date
    version_para = doc.add_paragraph()
    version_para.add_run('Version: 1.0\n').bold = True
    version_para.add_run('Generated: November 2025\n')
    version_para.add_run('Framework: Django Web Application\n')
    version_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Table of Contents
    doc.add_page_break()
    toc_heading = doc.add_heading('Table of Contents', 1)

    toc_items = [
        '1. Introduction',
        '2. System Requirements',
        '3. Installation and Setup',
        '4. Getting Started',
        '5. User Roles and Permissions',
        '6. Dashboard Overview',
        '7. Student Management',
        '8. Course Management',
        '9. Enrollment Management',
        '10. Grade Management',
        '11. Attendance Tracking',
        '12. Instructor Management',
        '13. Troubleshooting',
        '14. Technical Specifications',
        'Appendix A: API Endpoints',
        'Appendix B: Database Schema'
    ]

    for item in toc_items:
        doc.add_paragraph(item, style='List Number')

    # Section 1: Introduction
    doc.add_page_break()
    doc.add_heading('1. Introduction', 1)

    intro_text = """
    The Student Management System (SMS) is a comprehensive web-based application designed to streamline educational institution management. Built with Django and modern web technologies, it provides role-based access control for administrators, teachers, and students to manage academic data efficiently.

    Key Features:
    • Complete student lifecycle management
    • Course and curriculum administration
    • Enrollment and grade tracking
    • Attendance monitoring
    • Instructor management
    • Role-based permissions
    • Responsive web interface
    • File upload capabilities

    This manual provides detailed instructions for installing, configuring, and using all features of the Student Management System.
    """

    doc.add_paragraph(intro_text)

    # Section 2: System Requirements
    doc.add_heading('2. System Requirements', 1)

    req_table = doc.add_table(rows=1, cols=2)
    req_table.style = 'Table Grid'

    # Header
    hdr_cells = req_table.rows[0].cells
    hdr_cells[0].text = 'Component'
    hdr_cells[1].text = 'Requirement'

    # Requirements
    requirements = [
        ('Operating System', 'Windows 10/11, macOS, Linux'),
        ('Python Version', '3.8 or higher'),
        ('Database', 'MySQL 5.7+ (XAMPP recommended)'),
        ('Web Browser', 'Chrome, Firefox, Safari, Edge (latest versions)')
    ]

    for component, req in requirements:
        row_cells = req_table.add_row().cells
        row_cells[0].text = component
        row_cells[1].text = req

    # Section 3: Installation and Setup
    doc.add_page_break()
    doc.add_heading('3. Installation and Setup', 1)

    setup_steps = [
        ('Database Setup', '''1. Install XAMPP or MySQL Server
2. Start MySQL service
3. Create database: CREATE DATABASE student_management;
4. Note default credentials:
   • Host: localhost
   • Port: 3306
   • Username: root
   • Password: (empty)'''),

        ('Project Installation', '''1. Clone or download the project
2. Navigate to project directory
3. Create virtual environment (optional):
   python -m venv venv
   venv\\Scripts\\activate  # Windows
4. Install dependencies:
   pip install -r requirements.txt'''),

        ('Database Migration', '''1. Run migrations:
   python manage.py makemigrations
   python manage.py migrate
2. Load test data (optional):
   python manage.py shell < setup_test_data.py'''),

        ('Start Application', '''1. Run development server:
   python manage.py runserver
2. Open browser to: http://127.0.0.1:8000/
3. Login with provided credentials''')
    ]

    for title, content in setup_steps:
        doc.add_heading(title, 2)
        doc.add_paragraph(content)

    # Section 4: Getting Started
    doc.add_heading('4. Getting Started', 1)

    doc.add_heading('Login Credentials', 2)
    login_table = doc.add_table(rows=1, cols=3)
    login_table.style = 'Table Grid'

    # Header
    hdr_cells = login_table.rows[0].cells
    hdr_cells[0].text = 'Role'
    hdr_cells[1].text = 'Username'
    hdr_cells[2].text = 'Password'

    # Credentials
    credentials = [
        ('Administrator', 'admin', 'admin123'),
        ('Teacher', 'teacher1', 'teacher123'),
        ('Student', 'N/A', 'N/A (future feature)')
    ]

    for role, user, pwd in credentials:
        row_cells = login_table.add_row().cells
        row_cells[0].text = role
        row_cells[1].text = user
        row_cells[2].text = pwd

    doc.add_paragraph('\nNote: Use the test data setup script to create these accounts automatically.')

    # Section 5: User Roles and Permissions
    doc.add_page_break()
    doc.add_heading('5. User Roles and Permissions', 1)

    roles_data = [
        ('Administrator (Superuser)', '''Full system access including:
• Complete CRUD operations on all entities
• User and group management
• Django admin panel access
• System configuration
• All teacher permissions'''),

        ('Teacher', '''Limited administrative access:
• View students, courses, enrollments
• Add/edit/delete grades
• Add/edit/delete attendance records
• Cannot modify student or course records'''),

        ('Student (Future)', '''Self-service access:
• View personal information
• Access assigned courses
• Limited data visibility''')
    ]

    for role, permissions in roles_data:
        doc.add_heading(role, 2)
        doc.add_paragraph(permissions)

    # Section 6: Dashboard Overview
    doc.add_heading('6. Dashboard Overview', 1)

    dashboard_desc = """
    The dashboard serves as the main landing page after login, providing quick access to key information and actions.

    Key Elements:
    • Statistics cards showing total students and courses
    • Navigation sidebar with all main sections
    • Quick action buttons based on user permissions
    • Recent activity indicators
    • System status information
    """

    doc.add_paragraph(dashboard_desc)

    # Section 7: Student Management
    doc.add_page_break()
    doc.add_heading('7. Student Management', 1)

    student_features = [
        ('View Students', '''• Access student list with search functionality
• View detailed student profiles
• See enrollment history
• Access contact information'''),

        ('Add New Student (Admin Only)', '''• Fill required fields: Student ID, Name, Email, Date of Birth
• Add optional contact information
• Upload profile picture
• Automatic enrollment date assignment'''),

        ('Edit Student (Admin Only)', '''• Modify any student information
• Update or change profile picture
• Maintain data integrity'''),

        ('Delete Student (Admin Only)', '''• Remove student records
• Confirmation dialog prevents accidental deletion
• Cascading effects on related data''')
    ]

    for feature, desc in student_features:
        doc.add_heading(feature, 2)
        doc.add_paragraph(desc)

    # Section 8: Course Management
    doc.add_heading('8. Course Management', 1)

    course_features = [
        ('View Courses', '''• Browse all courses with search and filter options
• Filter by instructor
• View course details and enrolled students
• Access course information'''),

        ('Add New Course (Admin Only)', '''• Enter course code (unique identifier)
• Add course name and description
• Set credit hours
• Assign instructor from teacher group'''),

        ('Edit Course (Admin Only)', '''• Modify course information
• Change instructor assignment
• Update course details'''),

        ('Delete Course (Admin Only)', '''• Remove course records
• Confirmation required
• Check for existing enrollments''')
    ]

    for feature, desc in course_features:
        doc.add_heading(feature, 2)
        doc.add_paragraph(desc)

    # Section 9: Enrollment Management
    doc.add_page_break()
    doc.add_heading('9. Enrollment Management', 1)

    enrollment_desc = """
    Enrollment management handles the relationship between students and courses, ensuring proper course assignments and preventing duplicate enrollments.

    Key Features:
    • View all student-course enrollments
    • Add new enrollments (preventing duplicates)
    • Track enrollment dates
    • Link to grade and attendance records
    """

    doc.add_paragraph(enrollment_desc)

    # Section 10: Grade Management
    doc.add_heading('10. Grade Management', 1)

    grade_desc = """
    Grade management allows authorized users to assign and track academic performance with support for semester and year tracking.

    Features:
    • Assign grades A, B, C, D, F with color coding
    • Record semester and academic year
    • Link grades to specific enrollments
    • Permission-based access (Admin/Teacher only)
    """

    doc.add_paragraph(grade_desc)

    # Section 11: Attendance Tracking
    doc.add_heading('11. Attendance Tracking', 1)

    attendance_desc = """
    The attendance system provides daily tracking of student presence with date-based records.

    Features:
    • Record daily attendance (Present/Absent)
    • Date-based attendance entries
    • Prevent duplicate entries for same date
    • Link to enrollment records
    • Authorized user access only
    """

    doc.add_paragraph(attendance_desc)

    # Section 12: Instructor Management
    doc.add_page_break()
    doc.add_heading('12. Instructor Management', 1)

    instructor_desc = """
    Instructor management handles teacher accounts and their course assignments within the system.

    Features:
    • View all instructors and their course loads
    • Add new instructor accounts
    • Assign/remove teacher group membership
    • Track courses taught by each instructor
    """

    doc.add_paragraph(instructor_desc)

    # Section 13: Troubleshooting
    doc.add_heading('13. Troubleshooting', 1)

    issues = [
        ('Template Does Not Exist Error', '''Ensure all templates are in sms_app/templates/sms_app/ directory. Check for typos in template names.'''),

        ('Database Connection Error', '''• Verify MySQL server is running
• Check database credentials in settings.py
• Ensure database 'student_management' exists
• Confirm XAMPP MySQL service is active'''),

        ('Permission Denied Errors', '''• Check user roles and group assignments
• Verify @user_passes_test decorators
• Ensure proper authentication'''),

        ('File Upload Issues', '''• Check MEDIA_ROOT and MEDIA_URL settings
• Ensure proper permissions on media/ directory
• Verify file upload paths'''),

        ('Cannot Login', '''• Run setup_test_data.py to create users
• Check username/password combinations
• Verify user accounts exist in database''')
    ]

    for issue, solution in issues:
        doc.add_heading(issue, 2)
        doc.add_paragraph(solution)

    # Section 14: Technical Specifications
    doc.add_page_break()
    doc.add_heading('14. Technical Specifications', 1)

    tech_specs = [
        ('Backend Framework', 'Django 4.x'),
        ('Database', 'MySQL via XAMPP'),
        ('Frontend', 'HTML5, CSS3, Bootstrap 5'),
        ('Icons', 'Font Awesome'),
        ('Authentication', 'Django built-in auth system'),
        ('File Handling', 'Django media files'),
        ('Deployment', 'Development server (runserver)'),
        ('Python Version', '3.8+'),
        ('Operating System', 'Cross-platform (Windows/macOS/Linux)')
    ]

    tech_table = doc.add_table(rows=1, cols=2)
    tech_table.style = 'Table Grid'

    # Header
    hdr_cells = tech_table.rows[0].cells
    hdr_cells[0].text = 'Component'
    hdr_cells[1].text = 'Technology'

    for component, tech in tech_specs:
        row_cells = tech_table.add_row().cells
        row_cells[0].text = component
        row_cells[1].text = tech

    # Appendix A: API Endpoints
    doc.add_page_break()
    doc.add_heading('Appendix A: API Endpoints', 1)

    endpoints_table = doc.add_table(rows=1, cols=4)
    endpoints_table.style = 'Table Grid'

    # Header
    hdr_cells = endpoints_table.rows[0].cells
    hdr_cells[0].text = 'Endpoint'
    hdr_cells[1].text = 'Method'
    hdr_cells[2].text = 'Access'
    hdr_cells[3].text = 'Description'

    endpoints = [
        ('/', 'GET/POST', 'Public', 'Login page'),
        ('/dashboard/', 'GET', 'Auth', 'Main dashboard'),
        ('/students/', 'GET', 'Auth', 'Student list'),
        ('/students/add/', 'GET/POST', 'Admin', 'Add student'),
        ('/students/<id>/', 'GET', 'Auth', 'Student details'),
        ('/courses/', 'GET', 'Auth', 'Course list with search'),
        ('/courses/add/', 'GET/POST', 'Admin', 'Add course'),
        ('/enrollments/', 'GET', 'Auth', 'Enrollment list'),
        ('/grades/', 'GET', 'Auth', 'Grade list'),
        ('/grades/add/', 'GET/POST', 'Admin/Teacher', 'Add grade'),
        ('/attendance/', 'GET', 'Auth', 'Attendance records'),
        ('/attendance/add/', 'GET/POST', 'Admin/Teacher', 'Record attendance')
    ]

    for endpoint, method, access, desc in endpoints:
        row_cells = endpoints_table.add_row().cells
        row_cells[0].text = endpoint
        row_cells[1].text = method
        row_cells[2].text = access
        row_cells[3].text = desc

    # Appendix B: Database Schema
    doc.add_heading('Appendix B: Database Schema', 1)

    models_desc = [
        ('Student Model', '''• student_id: Unique identifier (CharField)
• first_name, last_name: Personal info (CharField)
• email: Contact (EmailField)
• phone, address: Optional contact (CharField/TextField)
• date_of_birth: Birth date (DateField)
• enrollment_date: Auto-generated (DateField)
• profile_picture: Optional image (ImageField)'''),

        ('Course Model', '''• course_code: Unique identifier (CharField)
• course_name: Display name (CharField)
• description: Course details (TextField)
• credits: Credit hours (IntegerField)
• instructor: Foreign key to User (ForeignKey)'''),

        ('Enrollment Model', '''• student: Foreign key to Student
• course: Foreign key to Course
• enrollment_date: Auto-generated timestamp
• Unique constraint: student-course pairs'''),

        ('Grade Model', '''• enrollment: Foreign key to Enrollment
• grade: A, B, C, D, F choices (CharField)
• semester: Semester name (CharField)
• academic_year: Format "2023-2024" (CharField)'''),

        ('Attendance Model', '''• enrollment: Foreign key to Enrollment
• date: Attendance date (DateField)
• status: Present/Absent choice (CharField)
• Unique constraint: enrollment-date pairs''')
    ]

    for model_name, fields in models_desc:
        doc.add_heading(model_name, 2)
        doc.add_paragraph(fields)

    # Save the document
    doc.save('doc-report.docx')
    print("User manual generated successfully: doc-report.docx")

if __name__ == '__main__':
    create_user_manual()